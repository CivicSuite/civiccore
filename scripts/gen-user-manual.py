"""Generate USER-MANUAL.docx and USER-MANUAL.pdf from USER-MANUAL.md.

Pandoc-free; uses python-docx + reportlab. Minimal markdown subset:
H1/H2/H3 headings, paragraphs, bullet lists, fenced code blocks, and images.
Generated DOCX/PDF outputs strip basic inline Markdown markers so the manual
reads as a document instead of a raw Markdown export.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as RLImage
from reportlab.platypus import Paragraph, Preformatted, SimpleDocTemplate, Spacer

IMAGE_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")
INLINE_CODE_RE = re.compile(r"`([^`]+)`")
BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "USER-MANUAL.md"
OUT_DOCX = ROOT / "USER-MANUAL.docx"
OUT_PDF = ROOT / "USER-MANUAL.pdf"


def _plain_inline(text: str) -> str:
    text = LINK_RE.sub(r"\1 (\2)", text)
    text = INLINE_CODE_RE.sub(r"\1", text)
    text = BOLD_RE.sub(r"\1", text)
    return text


def parse_blocks(md: str):
    """Yield (kind, text) blocks."""
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("# "):
            yield "h1", _plain_inline(line[2:].strip())
            i += 1
        elif line.startswith("## "):
            yield "h2", _plain_inline(line[3:].strip())
            i += 1
        elif line.startswith("### "):
            yield "h3", _plain_inline(line[4:].strip())
            i += 1
        elif line.startswith("```"):
            i += 1
            buf = []
            while i < len(lines) and not lines[i].startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            yield "code", "\n".join(buf)
        elif IMAGE_RE.match(line):
            match = IMAGE_RE.match(line)
            if match is not None:
                yield "image", f"{match.group(1)}|{match.group(2)}"
            i += 1
        elif line.lstrip().startswith(("- ", "* ")):
            yield "bullet", _plain_inline(line.lstrip()[2:].strip())
            i += 1
        elif line.strip() == "":
            i += 1
        else:
            yield "p", _plain_inline(line.strip())
            i += 1


def gen_docx(md: str, dest: Path) -> None:
    doc = Document()
    for kind, text in parse_blocks(md):
        if kind == "h1":
            doc.add_heading(text, level=1)
        elif kind == "h2":
            doc.add_heading(text, level=2)
        elif kind == "h3":
            doc.add_heading(text, level=3)
        elif kind == "code":
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif kind == "bullet":
            doc.add_paragraph(text, style="List Bullet")
        elif kind == "image":
            alt, path = text.split("|", 1)
            png_path = ROOT / path.replace(".svg", ".png")
            if png_path.exists():
                doc.add_picture(str(png_path), width=Inches(6))
            else:
                doc.add_paragraph(f"[image: {alt} ({path})]")
        else:
            doc.add_paragraph(text)
    doc.save(str(dest))


def _escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def gen_pdf(md: str, dest: Path) -> None:
    styles = getSampleStyleSheet()
    code_style = ParagraphStyle(
        "Code",
        parent=styles["Code"],
        fontName="Courier",
        fontSize=8,
        leftIndent=12,
        leading=10,
    )
    story = []
    for kind, text in parse_blocks(md):
        if kind == "h1":
            story.append(Paragraph(_escape(text), styles["Heading1"]))
        elif kind == "h2":
            story.append(Paragraph(_escape(text), styles["Heading2"]))
        elif kind == "h3":
            story.append(Paragraph(_escape(text), styles["Heading3"]))
        elif kind == "code":
            story.append(Preformatted(text, code_style))
        elif kind == "bullet":
            story.append(Paragraph("- " + _escape(text), styles["BodyText"]))
        elif kind == "image":
            alt, path = text.split("|", 1)
            png_path = ROOT / path.replace(".svg", ".png")
            if png_path.exists():
                from PIL import Image as PILImage

                with PILImage.open(str(png_path)) as image:
                    image_width, image_height = image.size
                target_width = 6.0 * inch
                target_height = target_width * (image_height / image_width)
                story.append(RLImage(str(png_path), width=target_width, height=target_height))
            else:
                story.append(
                    Paragraph(f"[image: {_escape(alt)} ({_escape(path)})]", styles["BodyText"])
                )
        else:
            story.append(Paragraph(_escape(text), styles["BodyText"]))
        story.append(Spacer(1, 4))

    doc = SimpleDocTemplate(
        str(dest),
        pagesize=LETTER,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    doc.build(story)


def main() -> None:
    md = SRC.read_text(encoding="utf-8")
    gen_docx(md, OUT_DOCX)
    gen_pdf(md, OUT_PDF)
    print(f"wrote {OUT_DOCX.relative_to(ROOT)} ({OUT_DOCX.stat().st_size} bytes)")
    print(f"wrote {OUT_PDF.relative_to(ROOT)} ({OUT_PDF.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
